# app.py
import streamlit as st
import re
import json
import os
from datetime import datetime
from typing import List, Dict, Any, Optional
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# Intenta cargar instaloader (para extraer captions de Instagram)
try:
    import instaloader
    HAS_INSTALOADER = True
except Exception:
    HAS_INSTALOADER = False

# ========== Configuración de página ==========
st.set_page_config(page_title="Recetario + Plan mensual", page_icon="🍽️", layout="wide")
st.title("📚 Recetario desde Instagram/TikTok + 📅 Plan mensual")

# ---------- Inicialización de estado ----------
def init_form_state():
    defaults = {
        "link": "",
        "caption_manual": "",
        "titulo": "",
        "porciones": "No especificado",
        "tiempo": "",
        "ingredientes_text": "",
        "procedimiento_text": "",
        "categoria": "Seleccionar opción",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_form_state()

# ========== Archivos ==========
RECETAS_FILE = "recetas.json"

def cargar_recetas(nombre_archivo: str = RECETAS_FILE) -> List[Dict[str, Any]]:
    if not os.path.exists(nombre_archivo):
        return []
    with open(nombre_archivo, "r", encoding="utf-8") as f:
        try:
            data = json.load(f)
            return data if isinstance(data, list) else []
        except json.JSONDecodeError:
            return []

def guardar_recetas(lista_recetas: List[Dict[str, Any]], nombre_archivo: str = RECETAS_FILE) -> None:
    with open(nombre_archivo, "w", encoding="utf-8") as f:
        json.dump(lista_recetas, f, ensure_ascii=False, indent=4)

# ========== Estilos DOCX ==========
def asegurar_estilos_docx(doc: Document):
    styles = doc.styles
    if 'Titulo1' not in styles:
        s1 = styles.add_style('Titulo1', WD_STYLE_TYPE.PARAGRAPH)
        s1.font.size = Pt(16); s1.font.bold = True
    if 'Titulo2' not in styles:
        s2 = styles.add_style('Titulo2', WD_STYLE_TYPE.PARAGRAPH)
        s2.font.size = Pt(14); s2.font.bold = True
    if 'Titulo3' not in styles:
        s3 = styles.add_style('Titulo3', WD_STYLE_TYPE.PARAGRAPH)
        s3.font.size = Pt(12); s3.font.bold = True

# ========== Helpers ==========
def capitalizar_oracion(texto: str) -> str:
    return texto[0].upper() + texto[1:] if texto else texto

def clean_bullet(text: str) -> str:
    return re.sub(r"^[\-\•\●\·\*✻❖❇️▪️✍🏼👨‍🍳]+\s*", "", text).strip()

def ig_shortcode_from_url(url: str) -> Optional[str]:
    try:
        parts = [p for p in url.split("/") if p.strip()]
        for p in reversed(parts):
            if re.fullmatch(r"[A-Za-z0-9_-]{5,20}", p):
                return p
        return None
    except Exception:
        return None

def get_instagram_caption(url: str) -> str:
    if HAS_INSTALOADER:
        try:
            L = instaloader.Instaloader(
                download_pictures=False, download_videos=False,
                download_video_thumbnails=False, download_comments=False,
                save_metadata=False, compress_json=False
            )
            shortcode = ig_shortcode_from_url(url)
            if shortcode:
                post = instaloader.Post.from_shortcode(L.context, shortcode)
                if post and post.caption:
                    return post.caption
        except Exception:
            pass
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        resp = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(resp.text, "html.parser")
        meta = soup.find("meta", attrs={"property": "og:description"}) or soup.find("meta", attrs={"name":"description"})
        if meta and meta.get("content"):
            return meta["content"]
    except Exception:
        pass
    return ""

def parse_recipe_from_caption(caption: str) -> Dict[str, Any]:
    rec = {"titulo": "", "porciones": "", "tiempo": "", "ingredientes": [], "procedimiento": []}
    if not caption:
        return rec

    lines = [l.strip() for l in caption.split("\n") if l.strip()]
    rec["titulo"] = lines[0] if lines else ""
    m_serves = re.search(r"(Serves|Porciones|Rinde)\s*[:\-]?\s*([0-9]+)", caption, flags=re.IGNORECASE)
    if m_serves:
        rec["porciones"] = m_serves.group(2).strip()
    m_time = re.search(r"(Takes|Tiempo)\s*[:\-]?\s*([0-9]+\s*\w+)", caption, flags=re.IGNORECASE)
    if m_time:
        rec["tiempo"] = m_time.group(2).strip()
    ing_split = re.split(r"(Ingredientes?:|Ingredients?:|✍🏼Ingredientes|👨‍🍳INGREDIENTES|❶ 𝐈𝐧𝐠𝐫𝐞𝐝𝐢𝐞𝐧𝐭𝐞𝐬:)", caption, flags=re.IGNORECASE)
    if len(ing_split) >= 2:
        after_ing = "".join(ing_split[1:])
        before_method = re.split(r"(Preparaci[oó]n:|Procedimiento:|Método:|Method:|❷ 𝐏𝐫𝐞𝐩𝐚𝐫𝐚𝐜𝐢[oó]n:)", after_ing, flags=re.IGNORECASE)[0]
        rec["ingredientes"] = [clean_bullet(x) for x in before_method.split("\n") if x.strip()]
        method_part = re.split(r"(Preparaci[oó]n:|Procedimiento:|Método:|Method:|❷ 𝐏𝐫𝐞𝐩𝐚𝐫𝐚𝐜𝐢[oó]n:)", after_ing, flags=re.IGNORECASE)
        if len(method_part) >= 2:
            rec["procedimiento"] = [clean_bullet(x) for x in "".join(method_part[1:]).split("\n") if x.strip()]
    if not rec["ingredientes"]:
        posibles_ingredientes = []
        posibles_procedimiento = []
        for l in lines[1:]:
            if re.match(r"^(\d+\/?\d*\s?(g|kg|ml|l|cucharad[ao]s?|tazas?|cdas?|cdtas?|pizca)?\s+.+)", l, flags=re.IGNORECASE):
                posibles_ingredientes.append(clean_bullet(l))
            else:
                posibles_procedimiento.append(clean_bullet(l))
        rec["ingredientes"] = posibles_ingredientes
        rec["procedimiento"] = posibles_procedimiento
    return rec

# ========== Sidebar ==========
pestanas = st.sidebar.radio(
    "Navegación",
    ["Nueva receta", "Ver recetas", "Exportar recetas", "Plan mensual"]
)

# ========== Pestaña: Nueva receta ==========
if pestanas == "Nueva receta":
    st.header("Agregar nueva receta desde enlace (Instagram/TikTok)")
    st.text_input("Ingresa el link del post:", key="link")
    col_a, col_b, col_c = st.columns([1,1,1])
    with col_a:
        if st.button("Leer descripción del enlace"):
            link_val = st.session_state.get("link", "").strip()
            if link_val:
                caption = get_instagram_caption(link_val) if "instagram.com" in link_val else ""
                if caption:
                    st.session_state.caption_manual = caption
                    parsed = parse_recipe_from_caption(caption)
                    st.session_state.titulo = parsed.get("titulo", "")
                    st.session_state.porciones = parsed.get("porciones", "No especificado")
                    st.session_state.tiempo = parsed.get("tiempo", "")
                    st.session_state.ingredientes_text = "\n".join(parsed.get("ingredientes", []))
                    st.session_state.procedimiento_text = "\n".join(parsed.get("procedimiento", []))
                    st.success("Descripción leída y campos actualizados.")
                else:
                    st.warning("No se pudo leer la descripción.")
    with col_b:
        if st.button("Rellenar desde el texto"):
            cap = st.session_state.get("caption_manual", "")
            if cap.strip():
                parsed = parse_recipe_from_caption(cap)
                st.session_state.titulo = parsed.get("titulo", "")
                st.session_state.porciones = parsed.get("porciones", "No especificado")
                st.session_state.tiempo = parsed.get("tiempo", "")
                st.session_state.ingredientes_text = "\n".join(parsed.get("ingredientes", []))
                st.session_state.procedimiento_text = "\n".join(parsed.get("procedimiento", []))
                st.success("Campos actualizados desde el texto.")
    with col_c:
        if st.button("Limpiar formulario"):
            init_form_state()
            st.info("Formulario limpio.")
    st.text_area("Descripción / receta:", key="caption_manual", height=200)
    st.subheader("📌 Datos de la receta")
    st.text_input("Nombre de la receta:", key="titulo")
    st.text_input("Porciones:", key="porciones")
    st.text_input("Tiempo:", key="tiempo")
    categorias = ["Sopa", "Proteína", "Arroz", "Guarnición", "Postre"]
    st.selectbox("Categoría:", ["Seleccionar opción"] + categorias, key="categoria")
    col1, col2 = st.columns(2)
    with col1:
        st.text_area("Ingredientes:", key="ingredientes_text", height=200)
    with col2:
        st.text_area("Procedimiento:", key="procedimiento_text", height=200)
    if st.button("Guardar receta"):
        if st.session_state.categoria == "Seleccionar opción":
            st.error("❌ Selecciona una categoría válida.")
        elif not st.session_state.titulo.strip():
            st.error("❌ Ingresa un nombre para la receta.")
        else:
            recetas = cargar_recetas()
            nueva = {
                "fuente": st.session_state.link.strip(),
                "titulo": capitalizar_oracion(st.session_state.titulo.strip()),
                "categoria": st.session_state.categoria.strip(),
                "porciones": st.session_state.porciones.strip() or "No especificado",
                "tiempo": st.session_state.tiempo.strip(),
                "ingredientes": [i.strip() for i in st.session_state.ingredientes_text.split("\n") if i.strip()],
                "procedimiento": [p.strip() for p in st.session_state.procedimiento_text.split("\n") if p.strip()],
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            recetas.append(nueva)
            guardar_recetas(recetas)
            st.success("✅ Receta guardada. Los datos se mantienen en el formulario.")

# ========== Pestaña: Ver recetas ==========
if pestanas == "Ver recetas":
    st.header("📖 Ver recetas guardadas")
    recetas = cargar_recetas()
    categorias = ["Sopa", "Proteína", "Arroz", "Guarnición", "Postre"]
    for cat in categorias:
        with st.expander(f"📂 {cat} ({len([r for r in recetas if r['categoria']==cat])} recetas)", expanded=False):
            recetas_cat = [r for r in recetas if r["categoria"] == cat]
            if recetas_cat:
                for receta in recetas_cat:
                    with st.expander(f"🍴 {receta['titulo']}", expanded=False):
                        st.text_input("Porciones:", value=receta["porciones"], key=f"{cat}_{receta['titulo']}_porciones")
                        st.text_input("Tiempo:", value=receta["tiempo"], key=f"{cat}_{receta['titulo']}_tiempo")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.text_area("Ingredientes:", value="\n".join(receta["ingredientes"]), height=150, key=f"{cat}_{receta['titulo']}_ing")
                        with col2:
                            st.text_area("Procedimiento:", value="\n".join(receta["procedimiento"]), height=150, key=f"{cat}_{receta['titulo']}_proc")
                        col_a, col_b = st.columns([1,1])
                        with col_a:
                            if st.button("🗑️ Eliminar receta", key=f"{cat}_{receta['titulo']}_del"):
                                recetas = [r for r in recetas if r["titulo"] != receta["titulo"] or r["categoria"] != receta["categoria"]]
                                guardar_recetas(recetas)
                                st.success("Receta eliminada.")
                                st.experimental_rerun()
                        with col_b:
                            if st.button("✏️ Editar receta", key=f"{cat}_{receta['titulo']}_edit"):
                                st.session_state.titulo = receta["titulo"]
                                st.session_state.porciones = receta["porciones"]
                                st.session_state.tiempo = receta["tiempo"]
                                st.session_state.ingredientes_text = "\n".join(receta["ingredientes"])
                                st.session_state.procedimiento_text = "\n".join(receta["procedimiento"])
                                st.session_state.link = receta.get("fuente", "")
                                st.session_state.categoria = receta["categoria"]
                                st.info("Edita los campos en la pestaña 'Nueva receta' y guarda para actualizar.")

# ========== Pestaña: Exportar recetas ==========
if pestanas == "Exportar recetas":
    st.header("📤 Exportar recetas a DOCX")
    recetas = cargar_recetas()
    if not recetas:
        st.info("No hay recetas para exportar.")
    else:
        doc = Document()
        asegurar_estilos_docx(doc)
        for receta in recetas:
            doc.add_paragraph(receta["titulo"], style="Titulo1")
            doc.add_paragraph(f"Categoría: {receta['categoria']}")
            doc.add_paragraph(f"Porciones: {receta['porciones']} | Tiempo: {receta['tiempo']}")
            doc.add_paragraph("Ingredientes:", style="Titulo2")
            for ing in receta["ingredientes"]:
                doc.add_paragraph(f"- {ing}")
            doc.add_paragraph("Procedimiento:", style="Titulo2")
            for step in receta["procedimiento"]:
                doc.add_paragraph(f"- {step}")
            doc.add_paragraph("\n")
        file_name = f"recetario_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(file_name)
        with open(file_name, "rb") as f:
            st.download_button("💾 Descargar DOCX", f, file_name=file_name)

# ========== Pestaña: Plan mensual ==========
if pestanas == "Plan mensual":
    st.header("🗓️ Generador de plan mensual")
    recetas = cargar_recetas()
    if not recetas:
        st.info("No hay recetas disponibles para el plan mensual.")
    else:
        dias_semana = ["Lunes","Martes","Miércoles","Jueves","Viernes","Sábado","Domingo"]
        seleccion = {}
        for dia in dias_semana:
            seleccion[dia] = st.selectbox(f"{dia}:", [""] + [r["titulo"] for r in recetas], key=f"plan_{dia}")
        if st.button("💾 Exportar plan mensual a DOCX"):
            doc = Document()
            asegurar_estilos_docx(doc)
            doc.add_paragraph("Plan mensual de recetas", style="Titulo1")
            for dia in dias_semana:
                receta_sel = next((r for r in recetas if r["titulo"] == seleccion[dia]), None)
                doc.add_paragraph(dia, style="Titulo2")
                if receta_sel:
                    doc.add_paragraph(receta_sel["titulo"])
                    doc.add_paragraph(f"Categoría: {receta_sel['categoria']}")
                    doc.add_paragraph(f"Porciones: {receta_sel['porciones']} | Tiempo: {receta_sel['tiempo']}")
                    doc.add_paragraph("Ingredientes:", style="Titulo3")
                    for ing in receta_sel["ingredientes"]:
                        doc.add_paragraph(f"- {ing}")
                    doc.add_paragraph("Procedimiento:", style="Titulo3")
                    for step in receta_sel["procedimiento"]:
                        doc.add_paragraph(f"- {step}")
                else:
                    doc.add_paragraph("No asignado")
            file_name = f"plan_mensual_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(file_name)
            with open(file_name, "rb") as f:
                st.download_button("💾 Descargar DOCX del plan", f, file_name=file_name)
